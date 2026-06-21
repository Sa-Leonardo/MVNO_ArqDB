from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Relatorio_Tecnico_MVNO.docx"
ASSETS = ROOT / "assets_relatorio"
ASSETS.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "18324A"
GRAY = "5E6873"
LIGHT_GRAY = "F2F4F7"
BORDER = "C8D0D9"
WHITE = "FFFFFF"
BLACK = "111827"
USABLE_DXA = 9360


def set_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == USABLE_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run("RELATÓRIO TÉCNICO | SISTEMA MVNO")
        set_font(header_run, size=8.5, color=GRAY, bold=True)
        add_page_number(sec.footer.paragraphs[0])


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run("[NOME DA INSTITUIÇÃO]"), size=12, color=GRAY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    set_font(p.add_run("[CURSO / TURMA]"), size=11, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("RELATÓRIO TÉCNICO"), size=14, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("ARQUITETURA E DESEMPENHO DE BANCO DE DADOS"), size=25, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(85)
    set_font(p.add_run("Sistema de Gestão MVNO com MongoDB"), size=15, color=GRAY)

    for label in ("Autor(a): [NOME DO(A) ALUNO(A)]", "Professor(a): [NOME DO(A) PROFESSOR(A)]", "Disciplina: Arquitetura e Desempenho de Banco de Dados"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(label), size=11, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(85)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("[CIDADE]"), size=11, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("2026"), size=11, color=GRAY)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold_prefix=None, justified=True):
    p = doc.add_paragraph()
    if justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        set_font(p.add_run(item))


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        set_font(p.add_run(item))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=9, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(str(value)), size=8.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7FA")
    p_pr.append(shd)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        set_font(p.add_run(line), name="Consolas", size=8.5, color="243447")
    return p


def make_placeholder(number, title):
    path = ASSETS / f"figura_{number:02d}.png"
    img = Image.new("RGB", (1500, 620), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 38)
        font_body = ImageFont.truetype("arial.ttf", 28)
    except OSError:
        font_title = ImageFont.load_default()
        font_body = ImageFont.load_default()
    for offset in range(0, 10, 2):
        draw.rectangle((30 + offset, 30 + offset, 1470 - offset, 590 - offset), outline="#AAB6C3", width=2)
    draw.text((750, 240), f"ESPAÇO RESERVADO PARA A FIGURA {number}", fill="#536273", anchor="mm", font=font_title)
    draw.text((750, 315), title, fill="#667788", anchor="mm", font=font_body)
    draw.text((750, 375), "Substitua este quadro pela imagem indicada", fill="#8A97A5", anchor="mm", font=font_body)
    img.save(path)
    return path


def add_figure_placeholder(doc, number, title, source="Fonte: elaboração própria / captura do sistema."):
    path = make_placeholder(number, title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(6.2))
    shape._inline.docPr.set("name", f"Figura {number} - espaço reservado")
    shape._inline.docPr.set("descr", f"Espaço reservado para inserir: {title}.")
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap.add_run(f"Figura {number} - {title}. "), size=9, color=GRAY, italic=True)
    set_font(cap.add_run(source), size=9, color=GRAY, italic=True)


def add_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EDF3F8")
    p_pr.append(shd)
    set_font(p.add_run(f"{title}: "), color=DARK_BLUE, bold=True)
    set_font(p.add_run(text), color="334155")


def add_summary(doc):
    add_heading(doc, "Resumo executivo", 1)
    add_paragraph(doc, "Este relatório apresenta a análise técnica do sistema de gestão de uma operadora móvel virtual (MVNO), desenvolvido como aplicação web para administrar clientes, chips, planos, assinaturas, recargas, lotes e usuários. O estudo concentra-se na arquitetura de software e, principalmente, nas decisões de modelagem, indexação, consistência e desempenho do MongoDB. A solução utiliza frontend React com TypeScript, API REST em Go com Gin e persistência documental no MongoDB. O backend segue separação em handlers, serviços, repositórios e modelos de domínio, com autenticação JWT e controle de acesso por papéis.")
    add_paragraph(doc, "A análise constatou uso adequado de documentos incorporados para informações com leitura conjunta, como contato e endereço do cliente, e uso de referências para relações de maior cardinalidade, como chips vinculados a clientes e assinaturas históricas. O sistema mantém snapshots de plano dentro de assinaturas para preservar o valor contratado no tempo. Foram identificados 17 índices explícitos, incluindo unicidade de ICCID, documento, nome de plano e e-mail de usuário. Como oportunidades de evolução, destacam-se paginação no servidor, transações nas operações multicoletânea, índice para lotes, testes automatizados, métricas reais no dashboard, validação de esquema e execução formal de benchmarks com explain('executionStats').")
    p = doc.add_paragraph()
    set_font(p.add_run("Palavras-chave: "), bold=True)
    set_font(p.add_run("MongoDB; banco de dados documental; MVNO; índices; desempenho; arquitetura em camadas; Go; React."))

    add_heading(doc, "Identificação do artefato analisado", 2)
    add_table(doc, ["Item", "Valor"], [
        ("Projeto", "MVNO_ArqDB - sistema administrativo para operação MVNO"),
        ("Repositório", "github.com/Sa-Leonardo/MVNO_ArqDB"),
        ("Branch analisada", "feat--implementacao-do-frontend"),
        ("Commit de referência", "b1c9bcb, acrescido das alterações locais descritas neste relatório"),
        ("Data da análise", "20 de junho de 2026"),
        ("Dimensão aproximada", "20 arquivos Go / 2.302 linhas; 38 arquivos TypeScript/TSX / 3.007 linhas"),
        ("Banco", "MongoDB, database padrão mvno_db"),
    ], [2700, 6660])
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Sumário", 1)
    items = [
        "1. Introdução",
        "2. Visão geral e requisitos do sistema",
        "3. Arquitetura de software",
        "4. Arquitetura de dados no MongoDB",
        "5. Modelo lógico das coleções",
        "6. Estratégia de índices",
        "7. Fluxos transacionais e consultas",
        "8. Análise de desempenho",
        "9. Plano de testes e benchmark",
        "10. Segurança, auditoria e privacidade",
        "11. Confiabilidade e concorrência",
        "12. Qualidade, testes e observabilidade",
        "13. Implantação e operação",
        "14. Recomendações priorizadas",
        "15. Conclusão",
        "Referências",
        "Apêndices",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        set_font(p.add_run(item), size=11)
    add_callout(doc, "Nota de edição", "No Microsoft Word, este sumário pode ser substituído por um sumário automático após a definição final das imagens e da paginação.")
    doc.add_page_break()


def chapter_1(doc):
    add_heading(doc, "1. Introdução", 1)
    add_heading(doc, "1.1 Contextualização", 2)
    add_paragraph(doc, "Uma MVNO (Mobile Virtual Network Operator) comercializa serviços móveis sem necessariamente possuir toda a infraestrutura física de radiofrequência. Sua operação depende de dados confiáveis sobre estoque de SIM cards, vínculo com clientes, planos contratados, histórico de assinaturas, recargas e ações administrativas. A qualidade da arquitetura de dados afeta diretamente a rastreabilidade, a velocidade de atendimento e a capacidade de crescimento da operação.")
    add_paragraph(doc, "O software analisado atende esse cenário por meio de uma aplicação web administrativa. O MongoDB foi escolhido como repositório principal e permite representar informações compostas em documentos BSON, preservando flexibilidade para campos opcionais e favorecendo leituras agregadas de entidades que costumam ser consultadas em conjunto.")
    add_heading(doc, "1.2 Problema", 2)
    add_paragraph(doc, "O problema técnico consiste em organizar dados de clientes, chips e produtos comerciais com integridade suficiente para impedir duplicidades e vínculos inválidos, ao mesmo tempo em que se mantêm consultas rápidas para telas operacionais. O sistema também deve registrar histórico, controlar permissões e suportar operações em lote sem comprometer consistência ou desempenho.")
    add_heading(doc, "1.3 Objetivos", 2)
    add_bullets(doc, [
        "Documentar a arquitetura atual do software e suas responsabilidades.",
        "Descrever o modelo documental, as coleções e os relacionamentos.",
        "Relacionar os índices existentes às consultas executadas pela aplicação.",
        "Avaliar riscos de desempenho, consistência, segurança e escalabilidade.",
        "Definir um plano reproduzível de testes de carga e análise com explain.",
        "Propor melhorias priorizadas para evolução acadêmica e produtiva.",
    ])
    add_heading(doc, "1.4 Escopo e metodologia", 2)
    add_paragraph(doc, "A metodologia empregada foi análise estática do código-fonte, inspeção das dependências, rotas, modelos BSON, métodos de repositório e índices criados na inicialização. Também foram executadas as verificações go test ./... e npm run build. Como não foi fornecida uma massa controlada de produção nem um ambiente dedicado de carga, este relatório separa fatos observados de hipóteses de desempenho e fornece quadros próprios para inserir resultados experimentais.")
    add_callout(doc, "Delimitação", "Valores de latência, throughput e uso de recursos não são inventados. Eles devem ser preenchidos após a execução do protocolo do Capítulo 9 no ambiente da disciplina.")


def chapter_2(doc):
    add_heading(doc, "2. Visão geral e requisitos do sistema", 1)
    add_heading(doc, "2.1 Módulos funcionais", 2)
    rows = [
        ("Autenticação", "Login com e-mail e senha; emissão e validação de JWT."),
        ("Usuários", "Cadastro, consulta, edição e desativação, restritos ao administrador."),
        ("Clientes", "Cadastro com ao menos um chip, edição, listagem e exclusão condicionada."),
        ("Chips", "Cadastro individual ou em lote, busca, filtro por status e detalhamento."),
        ("Ativação", "Vínculo entre chip, cliente e plano com criação da assinatura."),
        ("Planos", "Cadastro, edição, ativação lógica, preço, ciclo e benefícios."),
        ("Recargas", "Criação e histórico por ICCID."),
        ("Dashboard", "Indicadores de clientes e chips; parte dos dados financeiros é simulada."),
        ("Auditoria", "Registro de ações com entidade, ator, instante e snapshot."),
    ]
    add_table(doc, ["Módulo", "Responsabilidade"], rows, [2300, 7060])
    add_heading(doc, "2.2 Perfis e autorização", 2)
    add_table(doc, ["Perfil", "Leitura", "Operação", "Administração"], [
        ("admin", "Todos os módulos", "Clientes, chips, lotes, planos e recargas", "Usuários e exclusão de clientes"),
        ("operator", "Dados operacionais", "Clientes, chips, lotes, ativações e recargas", "Sem gestão de usuários"),
        ("viewer", "Consultas protegidas", "Nenhuma escrita", "Nenhuma"),
    ], [1500, 2450, 3000, 2410])
    add_heading(doc, "2.3 Regras de negócio relevantes", 2)
    add_bullets(doc, [
        "Cada cliente novo deve receber pelo menos um chip disponível.",
        "O documento do cliente, o ICCID do chip, o nome do plano e o e-mail do usuário são únicos.",
        "Um chip reservado só pode ser ativado para o cliente ao qual foi previamente associado.",
        "A exclusão de cliente é bloqueada quando existem chips vinculados.",
        "A criação de lote gera de 1 a 1.000 chips sequenciais e rejeita ICCID já cadastrado.",
        "O histórico de assinatura armazena snapshot do plano para preservar informação temporal.",
    ])
    add_figure_placeholder(doc, 1, "Diagrama de casos de uso do sistema")


def chapter_3(doc):
    add_heading(doc, "3. Arquitetura de software", 1)
    add_heading(doc, "3.1 Visão de contexto", 2)
    add_paragraph(doc, "O usuário acessa uma SPA executada no navegador. O frontend chama a API REST por HTTP e envia o JWT no cabeçalho Authorization. A API valida a identidade e o papel, executa as regras de negócio e usa o driver oficial do MongoDB para ler e alterar documentos. No desenvolvimento, o Vite encaminha /api, /auth e /health para localhost:8080.")
    add_figure_placeholder(doc, 2, "Diagrama de contexto: usuário, frontend, API e MongoDB")
    add_heading(doc, "3.2 Tecnologias", 2)
    add_table(doc, ["Camada", "Tecnologia", "Função"], [
        ("Apresentação", "React 18, TypeScript, Vite, Tailwind CSS", "SPA administrativa responsiva."),
        ("Estado e HTTP", "TanStack Query, Axios", "Cache de requisições, mutações e consumo REST."),
        ("Visualização", "Recharts, Lucide React", "Gráficos e iconografia."),
        ("API", "Go 1.25, Gin", "Rotas, middleware, validação e respostas HTTP."),
        ("Persistência", "MongoDB Driver for Go v1.17", "CRUD, índices e operações em lote."),
        ("Segurança", "JWT v5 e bcrypt", "Sessão stateless e hash de senhas."),
        ("Observabilidade", "Zap e gin.Logger", "Logs estruturados e de requisição."),
    ], [1800, 3000, 4560])
    add_heading(doc, "3.3 Arquitetura em camadas", 2)
    add_numbered(doc, [
        "Handler: recebe HTTP, faz binding dos DTOs, traduz erros e retorna JSON.",
        "Service: concentra regras de negócio, validações e composição de operações.",
        "Repository: encapsula filtros BSON, índices, ordenação e persistência.",
        "Model: define estruturas de domínio e mapeamento BSON/JSON.",
        "Database/Config: inicializa conexão, variáveis de ambiente e ciclo de vida.",
    ])
    add_figure_placeholder(doc, 3, "Diagrama de componentes e dependências internas")
    add_heading(doc, "3.4 Fluxo de uma requisição", 2)
    add_paragraph(doc, "Uma chamada protegida passa pelo middleware Auth, que extrai o token Bearer, valida assinatura e expiração e injeta user_id, e-mail e papel no contexto Gin. Quando necessário, RequireRole limita os papéis permitidos. O handler converte o JSON em DTO, o serviço aplica regras, o repositório acessa o MongoDB e a resposta é encapsulada em um formato uniforme.")
    add_figure_placeholder(doc, 4, "Diagrama de sequência da ativação de um chip")


def chapter_4(doc):
    add_heading(doc, "4. Arquitetura de dados no MongoDB", 1)
    add_heading(doc, "4.1 Justificativa do modelo documental", 2)
    add_paragraph(doc, "O domínio combina entidades relativamente estáveis, estruturas internas pequenas e históricos crescentes. O modelo documental é apropriado para contato, endereço, benefícios e metadados porque esses dados pertencem à entidade principal e são normalmente carregados juntos. Para históricos e relações de cardinalidade elevada, o projeto usa coleções independentes e referências por ObjectID ou ICCID.")
    add_heading(doc, "4.2 Coleções", 2)
    collections = [
        ("users", "Usuários, credenciais, papel, status e metadados", "Baixa"),
        ("clientes", "Dados cadastrais e contato", "Média"),
        ("chips", "Estoque e estado operacional de SIM cards", "Alta"),
        ("planos", "Catálogo comercial", "Baixa"),
        ("assinaturas", "Histórico de contratação por chip", "Alta e crescente"),
        ("recargas", "Eventos financeiros por chip", "Alta e crescente"),
        ("audit_logs", "Eventos de auditoria", "Alta e crescente"),
        ("lotes", "Metadados de importação em lote", "Média"),
    ]
    add_table(doc, ["Coleção", "Conteúdo", "Crescimento esperado"], collections, [1900, 5000, 2460])
    add_heading(doc, "4.3 Incorporação e referências", 2)
    add_table(doc, ["Decisão", "Exemplos", "Motivação"], [
        ("Incorporar", "cliente.contato, cliente.endereco, plano.beneficios", "Leitura conjunta, tamanho limitado e mesma vida útil."),
        ("Incorporar snapshot", "assinatura.plano e chip.assinatura_atual", "Evita junção na leitura e preserva valor histórico."),
        ("Referenciar", "chip.cliente_id, chip.plano_id, chip.lote_id", "Entidades independentes e relacionamento consultável."),
        ("Separar histórico", "assinaturas, recargas e audit_logs", "Crescimento sem ampliar indefinidamente o documento do chip."),
    ], [1800, 3100, 4460])
    add_heading(doc, "4.4 Relacionamentos e cardinalidade", 2)
    add_bullets(doc, [
        "Cliente 1:N Chip, materializado em chips.cliente_id.",
        "Lote 1:N Chip, materializado em chips.lote_id.",
        "Plano 1:N Assinatura, preservado no snapshot plano.plano_id.",
        "Chip 1:N Assinatura, relacionado por assinaturas.iccid.",
        "Chip 1:N Recarga, relacionado por recargas.iccid.",
        "Usuário 1:N Evento de auditoria, relacionado por audit_logs.actor.user_id.",
    ])
    add_figure_placeholder(doc, 5, "Modelo lógico documental das coleções MongoDB")
    add_heading(doc, "4.5 Consistência e desnormalização", 2)
    add_paragraph(doc, "A desnormalização reduz leituras adicionais no detalhamento do chip, pois assinatura e plano atuais já estão incorporados. O custo é a necessidade de atualização coordenada quando o estado muda. O snapshot histórico não deve acompanhar alterações posteriores do plano, pois sua finalidade é registrar as condições vigentes na contratação.")
    add_callout(doc, "Risco identificado", "Ativação e criação de lote alteram mais de uma coleção sem transação MongoDB. Uma falha intermediária pode deixar assinatura ou lote sem o correspondente estado final do chip.")


def chapter_5(doc):
    add_heading(doc, "5. Modelo lógico das coleções", 1)
    add_heading(doc, "5.1 Clientes", 2)
    add_table(doc, ["Campo", "Tipo", "Descrição"], [
        ("_id", "ObjectId", "Identificador gerado pelo MongoDB."),
        ("nome", "string", "Nome ou razão social."),
        ("documento", "string", "CPF/CNPJ normalizado; único."),
        ("contato", "documento", "E-mail e telefone opcionais."),
        ("endereco", "documento", "Logradouro, número, cidade, UF e CEP."),
        ("tags", "array<string>", "Classificação multivalorada."),
        ("audit", "documento", "created_at e updated_at."),
    ], [2100, 1800, 5460])
    add_heading(doc, "5.2 Chips", 2)
    add_table(doc, ["Campo", "Tipo", "Descrição"], [
        ("iccid", "string", "Identificador físico único do SIM."),
        ("msisdn", "string", "Número telefônico, opcional."),
        ("status", "enum", "available, reserved, active, suspended ou canceled."),
        ("lote_id", "ObjectId?", "Referência ao lote de origem."),
        ("cliente_id", "ObjectId?", "Referência ao titular."),
        ("plano_id", "ObjectId?", "Referência ao plano atual."),
        ("rede", "documento", "Operadora e IMSI."),
        ("assinatura_atual", "documento?", "Snapshot otimizado para leitura operacional."),
        ("tags/audit", "array/documento", "Classificação e timestamps."),
    ], [2100, 1800, 5460])
    add_heading(doc, "5.3 Planos, assinaturas e recargas", 2)
    add_paragraph(doc, "Planos armazenam nome, descrição, valor, moeda, ciclo em dias, benefícios e flag ativo. Assinaturas registram ICCID, cliente, snapshot do plano, status e período de vigência. Recargas armazenam ICCID, valor, moeda, status, referência e timestamps de solicitação/processamento. A separação mantém os históricos extensíveis e permite índices orientados por chip e data.")
    add_heading(doc, "5.4 Usuários, auditoria e lotes", 2)
    add_paragraph(doc, "Usuários são documentos compostos por identity, credentials, access, status, audit e metadata. A senha nunca é serializada no JSON. Audit logs guardam entidade, identificador, ação, ator, instante e snapshot. Lotes registram nome, descrição, quantidade, data, status e auditoria; cada chip gerado carrega lote_id.")
    add_figure_placeholder(doc, 6, "Exemplo das coleções no MongoDB Compass")
    add_heading(doc, "5.5 Exemplo de documento de chip", 2)
    add_code(doc, '''{
  "_id": ObjectId("..."),
  "iccid": "89550001000001",
  "msisdn": "55859000001",
  "status": "active",
  "cliente_id": ObjectId("..."),
  "plano_id": ObjectId("..."),
  "rede": { "operadora": "MVNO", "imsi": "72400000001" },
  "assinatura_atual": {
    "assinatura_id": ObjectId("..."),
    "status": "active",
    "plano": { "nome": "Plano 20 GB", "valor": 49.90, "moeda": "BRL" }
  },
  "tags": ["lote", "turma-a"],
  "audit": { "created_at": ISODate("..."), "updated_at": ISODate("...") }
}''')


def chapter_6(doc):
    add_heading(doc, "6. Estratégia de índices", 1)
    add_heading(doc, "6.1 Índices implementados", 2)
    rows = [
        ("chips", "uniq_chips_iccid", "iccid ASC", "Único", "Busca e integridade por ICCID"),
        ("chips", "idx_chips_msisdn", "msisdn ASC", "Normal", "Busca por número"),
        ("chips", "idx_chips_status_cliente", "status, cliente_id", "Composto", "Filtros e vínculos"),
        ("chips", "idx_chips_assinatura_status", "assinatura_atual.status", "Normal", "Estado da assinatura"),
        ("clientes", "uniq_clientes_documento", "documento ASC", "Único", "Integridade cadastral"),
        ("clientes", "idx_clientes_nome", "nome ASC", "Normal", "Ordenação/pesquisa"),
        ("clientes", "idx_clientes_tags", "tags ASC", "Multikey", "Filtro por tag"),
        ("planos", "uniq_planos_nome", "nome ASC", "Único", "Integridade do catálogo"),
        ("planos", "idx_planos_ativo_valor", "ativo, valor", "Composto", "Lista de ativos por preço"),
        ("recargas", "idx_recargas_iccid_data", "iccid, solicitada_em DESC", "Composto", "Histórico por chip"),
        ("recargas", "idx_recargas_status", "status ASC", "Normal", "Processamento por status"),
        ("assinaturas", "idx_assinaturas_iccid_inicio", "iccid, inicio_em DESC", "Composto", "Histórico por chip"),
        ("assinaturas", "idx_assinaturas_cliente_status", "cliente_id, status", "Composto", "Carteira do cliente"),
        ("audit_logs", "idx_audit_entity", "entity, entity_id, occurred_at DESC", "Composto", "Trilha da entidade"),
        ("audit_logs", "idx_audit_actor", "actor.user_id, occurred_at DESC", "Composto", "Ações por usuário"),
        ("users", "uniq_users_identity_email", "identity.email ASC", "Único parcial", "Login e integridade"),
        ("users", "idx_users_status_role", "status.is_active, access.role", "Composto", "Administração"),
        ("users", "idx_users_metadata_tags", "metadata.tags ASC", "Multikey", "Filtro por tag"),
    ]
    add_table(doc, ["Coleção", "Índice", "Chaves", "Tipo", "Finalidade"], rows, [1300, 2300, 2400, 1300, 2060])
    add_heading(doc, "6.2 Adequação às consultas", 2)
    add_paragraph(doc, "Os índices compostos de recargas e assinaturas seguem a regra prática de igualdade antes da ordenação: primeiro filtram pelo ICCID e depois entregam os eventos em ordem temporal decrescente. O índice de planos combina o filtro ativo com a ordenação por valor. Os índices únicos funcionam simultaneamente como aceleradores de busca e restrições de integridade.")
    add_heading(doc, "6.3 Lacunas e custo de manutenção", 2)
    add_bullets(doc, [
        "ListChips ordena por audit.created_at, mas não existe índice com essa chave; em bases grandes pode ocorrer SORT em memória.",
        "ListLotes ordena por criado_em sem índice específico.",
        "CountChipsByClienteID pode usar o segundo campo do índice status, cliente_id apenas de forma limitada; um índice iniciado por cliente_id seria mais adequado.",
        "Índices de baixa seletividade isolada, como recargas.status, devem ser justificados por consultas reais.",
        "Cada índice aumenta armazenamento, memória do working set e custo de InsertOne/InsertMany/UpdateMany.",
    ])
    add_figure_placeholder(doc, 7, "Lista de índices no MongoDB Compass")
    add_heading(doc, "6.4 Índices recomendados para avaliação", 2)
    add_code(doc, '''db.chips.createIndex({ "audit.created_at": -1 }, { name: "idx_chips_created_desc" })
db.chips.createIndex({ cliente_id: 1 }, { name: "idx_chips_cliente" })
db.lotes.createIndex({ criado_em: -1 }, { name: "idx_lotes_criado_desc" })''')
    add_callout(doc, "Critério", "Nenhum índice novo deve ser adotado apenas por hipótese. Compare executionStats, latência e impacto de escrita antes e depois.")


def chapter_7(doc):
    add_heading(doc, "7. Fluxos transacionais e consultas", 1)
    add_heading(doc, "7.1 Cadastro de cliente com chips", 2)
    add_numbered(doc, [
        "Normalizar e remover ICCIDs repetidos da requisição.",
        "Consultar todos os chips por $in e validar status available e ausência de cliente.",
        "Inserir o cliente.",
        "Executar UpdateMany para reservar os chips e associar cliente_id.",
        "Se a associação falhar, excluir o cliente como compensação.",
        "Registrar auditoria de criação.",
    ])
    add_paragraph(doc, "A compensação reduz inconsistências, mas não oferece a mesma garantia de uma transação: concorrência ou falha entre etapas ainda pode produzir estados intermediários. Em implantação com replica set, uma sessão transacional pode tornar a inserção do cliente e a reserva dos chips atômicas.")
    add_heading(doc, "7.2 Ativação", 2)
    add_paragraph(doc, "O serviço carrega chip, cliente e plano, valida status e titularidade, cria uma assinatura com snapshot do plano e atualiza o chip. O repositório primeiro insere a assinatura e depois atualiza o chip. Essa ordem preserva o histórico, porém pode gerar assinatura sem chip ativo se a segunda escrita falhar.")
    add_heading(doc, "7.3 Cadastro em lote", 2)
    add_paragraph(doc, "O lote aceita até 1.000 itens. O backend gera sufixos de seis dígitos, consulta previamente todos os ICCIDs e usa InsertMany para reduzir round trips. O lote é inserido antes dos chips; portanto, uma falha na operação em massa pode deixar metadado de lote sem seus itens. A solução recomendada é transação ou status de processamento com retomada idempotente.")
    add_heading(doc, "7.4 Exclusão de cliente", 2)
    add_paragraph(doc, "A exclusão física é permitida apenas quando CountDocuments confirma ausência de chips vinculados. A regra evita referência órfã imediata. Para requisitos legais ou auditoria mais forte, seria preferível exclusão lógica com deleted_at, retenção e anonimização controlada.")
    add_heading(doc, "7.5 Padrões de consulta", 2)
    add_table(doc, ["Operação", "Filtro", "Ordenação", "Observação"], [
        ("Listar clientes", "{}", "nome ASC", "Retorna toda a coleção."),
        ("Listar planos", "ativo=true", "valor ASC", "Coberto por índice composto."),
        ("Listar chips", "status opcional", "audit.created_at DESC", "Sem paginação e sem índice da ordenação."),
        ("Histórico de recargas", "iccid", "solicitada_em DESC", "Bem alinhado ao índice."),
        ("Histórico de assinaturas", "iccid", "inicio_em DESC", "Bem alinhado ao índice."),
        ("Listar lotes", "{}", "criado_em DESC", "Sem índice específico."),
    ], [2100, 2100, 2600, 2560])


def chapter_8(doc):
    add_heading(doc, "8. Análise de desempenho", 1)
    add_heading(doc, "8.1 Características da carga", 2)
    add_paragraph(doc, "A carga operacional tende a ser majoritariamente de leitura nas telas de dashboard, listagens e detalhes, com picos de escrita em importação de chips, ativações e recargas. Assinaturas, recargas e auditoria são coleções append-heavy. Chips sofrem atualizações de estado. Clientes e planos têm menor taxa de alteração.")
    add_heading(doc, "8.2 Pontos positivos", 2)
    add_bullets(doc, [
        "Índices únicos eliminam duplicidades no nível do banco.",
        "InsertMany reduz chamadas no cadastro em lote.",
        "Snapshots diminuem leituras adicionais para exibir plano atual.",
        "Consultas históricas usam índices compostos compatíveis com filtro e ordenação.",
        "O driver mantém pool de conexões por cliente MongoDB compartilhado.",
        "Timeouts existem na conexão, health check, desconexão e shutdown HTTP.",
    ])
    add_heading(doc, "8.3 Gargalos potenciais", 2)
    add_table(doc, ["Risco", "Efeito", "Evidência no código", "Tratamento"], [
        ("Listagens sem paginação", "Memória, rede e tempo crescem linearmente", "cursor.All em clientes, chips, planos e usuários", "Cursor/limit e paginação por chave"),
        ("Dashboard no cliente", "Baixa todos os chips para contar status", "Filtros e contagens no React", "Endpoint agregado com $group"),
        ("Ordenação sem índice", "SORT em memória", "chips.audit.created_at e lotes.criado_em", "Índices após benchmark"),
        ("Auditoria crescente", "Working set e armazenamento", "Coleção append-only sem retenção", "Particionamento/arquivamento/TTL conforme regra"),
        ("Sem projeção", "Mais bytes lidos e enviados", "Find retorna documentos completos", "Projection por tela"),
        ("Sem limites de contexto por request", "Operações longas", "Contexto HTTP sem timeout de banco explícito", "Middleware e context.WithTimeout"),
    ], [1800, 2200, 2760, 2600])
    add_heading(doc, "8.4 Escalabilidade", 2)
    add_paragraph(doc, "Antes de considerar sharding, a prioridade é reduzir o volume por requisição, criar índices guiados por evidência e separar consultas analíticas. Em escala maior, recargas, assinaturas e auditoria são candidatas naturais à distribuição. A escolha da shard key exigiria análise de cardinalidade, frequência e monotonicidade; ICCID com componente hash pode distribuir eventos, mas afeta consultas por intervalo temporal. Essa decisão não deve ser antecipada sem métricas.")
    add_heading(doc, "8.5 Complexidade aproximada", 2)
    add_paragraph(doc, "Buscas por chaves indexadas têm custo aproximado logarítmico sobre a árvore B, além do fetch do documento. Sem índice adequado, uma consulta pode examinar N documentos. InsertMany reduz overhead de rede, mas ainda atualiza todos os índices por documento. Ordenações não cobertas exigem memória proporcional ao conjunto retornado e podem derramar para disco quando permitido.")


def chapter_9(doc):
    add_heading(doc, "9. Plano de testes e benchmark", 1)
    add_heading(doc, "9.1 Objetivo experimental", 2)
    add_paragraph(doc, "O benchmark deve medir a influência de volume, índices e concorrência sobre as consultas críticas. Recomenda-se executar em máquina isolada, repetir cada cenário ao menos cinco vezes, descartar aquecimento e registrar mediana e percentil 95. O ambiente, versões e limites de hardware devem permanecer constantes.")
    add_heading(doc, "9.2 Massa de dados", 2)
    add_table(doc, ["Cenário", "Clientes", "Chips", "Assinaturas", "Recargas", "Auditoria"], [
        ("Pequeno", "1.000", "5.000", "5.000", "25.000", "20.000"),
        ("Médio", "10.000", "100.000", "150.000", "1.000.000", "500.000"),
        ("Grande", "100.000", "1.000.000", "2.000.000", "10.000.000", "5.000.000"),
    ], [1500, 1500, 1500, 1700, 1600, 1560])
    add_heading(doc, "9.3 Consultas a medir", 2)
    add_numbered(doc, [
        "Busca exata de chip por ICCID.",
        "Listagem de chips por status, ordenada por criação, com e sem índice proposto.",
        "Histórico das últimas 50 recargas de um ICCID.",
        "Histórico das últimas 50 assinaturas de um ICCID.",
        "Contagem de chips por cliente.",
        "Listagem de lotes por data, com e sem índice.",
        "Importação de lotes com 10, 100 e 1.000 chips.",
        "Login concorrente e leitura do dashboard agregado proposto.",
    ])
    add_heading(doc, "9.4 Coleta com explain", 2)
    add_code(doc, '''db.chips.find({ status: "active" })
  .sort({ "audit.created_at": -1 })
  .limit(50)
  .explain("executionStats")

db.recargas.find({ iccid: "89550001000001" })
  .sort({ solicitada_em: -1 })
  .limit(50)
  .explain("executionStats")''')
    add_bullets(doc, [
        "executionTimeMillis: duração observada no servidor.",
        "totalKeysExamined: entradas de índice percorridas.",
        "totalDocsExamined: documentos lidos.",
        "nReturned: documentos efetivamente retornados.",
        "winningPlan: verificar IXSCAN/COLLSCAN e estágio SORT.",
        "Razão docsExamined/nReturned: indicador simples de eficiência.",
    ])
    add_figure_placeholder(doc, 8, "Explain antes da otimização")
    add_figure_placeholder(doc, 9, "Explain depois da otimização")
    add_heading(doc, "9.5 Quadro para resultados", 2)
    add_table(doc, ["Consulta", "Volume", "Índice", "P50 (ms)", "P95 (ms)", "Docs exam.", "Resultado"], [
        ("Busca ICCID", "[preencher]", "uniq_chips_iccid", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("Chips por status/data", "[preencher]", "antes", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("Chips por status/data", "[preencher]", "depois", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("Recargas por ICCID", "[preencher]", "idx_recargas_iccid_data", "[ ]", "[ ]", "[ ]", "[ ]"),
        ("InsertMany 1.000", "1.000", "índices atuais", "[ ]", "[ ]", "n/a", "[ ]"),
    ], [1800, 1300, 2050, 1000, 1000, 1100, 1110])
    add_figure_placeholder(doc, 10, "Gráfico comparativo de latência e throughput")
    add_heading(doc, "9.6 Critérios de aceitação sugeridos", 2)
    add_bullets(doc, [
        "Consultas pontuais devem usar IXSCAN e examinar quantidade próxima ao retorno.",
        "Listagens devem ter limite explícito e tempo previsível com crescimento da base.",
        "P95 deve ser reportado, pois a média pode esconder caudas de latência.",
        "Nenhum teste deve produzir duplicidade, lote parcial silencioso ou assinatura órfã.",
        "O ganho de leitura deve ser comparado ao aumento de custo de escrita e armazenamento.",
    ])


def chapter_10(doc):
    add_heading(doc, "10. Segurança, auditoria e privacidade", 1)
    add_heading(doc, "10.1 Controles existentes", 2)
    add_bullets(doc, [
        "Senhas armazenadas como hash bcrypt e omitidas das respostas JSON.",
        "JWT assinado com segredo e expiração configuráveis.",
        "Middleware Bearer para todas as rotas /api/v1.",
        "RBAC com admin, operator e viewer.",
        "DTOs com validações de obrigatoriedade, formato e limites.",
        "Auditoria de ações de domínio com ator e snapshot.",
        "Índices únicos como última barreira contra duplicidade.",
    ])
    add_heading(doc, "10.2 Riscos e melhorias", 2)
    add_table(doc, ["Tema", "Situação", "Recomendação"], [
        ("Segredo JWT", "Exemplo contém valor previsível", "Usar secret manager e falhar ao iniciar sem segredo forte."),
        ("Usuário desativado", "JWT já emitido continua válido até expirar", "Validar status ou usar revogação/versão de token."),
        ("Transporte", "Desenvolvimento usa HTTP", "TLS obrigatório em produção."),
        ("Rate limiting", "Não identificado", "Limitar login e operações de escrita."),
        ("Dados pessoais", "CPF/CNPJ, contato e endereço", "Política LGPD, minimização, retenção e controle de exportação."),
        ("Auditoria", "Snapshot pode carregar PII", "Reduzir campos sensíveis e definir retenção."),
        ("CORS/headers", "Não há política explícita", "Definir origens, CSP e headers de segurança."),
    ], [1900, 3200, 4260])
    add_heading(doc, "10.3 Matriz de ameaça resumida", 2)
    add_paragraph(doc, "Os principais vetores são roubo de token, força bruta no login, abuso de papel administrativo, exposição de dados pessoais, manipulação de identificadores e negação de serviço por listagens sem limite. A mitigação combina expiração curta, TLS, rate limiting, autorização no servidor, paginação, logs sem segredos, validação de entrada e monitoramento de falhas.")


def chapter_11(doc):
    add_heading(doc, "11. Confiabilidade e concorrência", 1)
    add_heading(doc, "11.1 Ciclo de vida", 2)
    add_paragraph(doc, "A aplicação testa o MongoDB na inicialização, cria índices antes de expor o serviço, oferece /health com ping de dois segundos e implementa shutdown gracioso de dez segundos. O cliente MongoDB é desconectado com timeout de cinco segundos. Esses mecanismos reduzem perda de requisições no encerramento e detectam indisponibilidade básica.")
    add_heading(doc, "11.2 Atomicidade", 2)
    add_paragraph(doc, "Atualizações de um único documento são atômicas no MongoDB. Entretanto, cadastro de cliente, ativação e importação em lote atravessam mais de um documento ou coleção. Para garantir invariantes fortes, recomenda-se executar essas etapas em transações quando o MongoDB estiver configurado como replica set. Alternativamente, usar estados pending/failed, chaves idempotentes e processos de reconciliação.")
    add_heading(doc, "11.3 Condições de corrida", 2)
    add_bullets(doc, [
        "Dois operadores podem tentar reservar o mesmo chip; o filtro atômico do UpdateMany reduz o risco e ModifiedCount detecta conflito.",
        "A validação prévia do lote e o InsertMany não são uma única operação lógica; o índice único ainda impede duplicidade, mas pode ocorrer lote parcial conforme ordered insert.",
        "A ativação lê o estado e escreve depois; a atualização deveria incluir status esperado no filtro ou usar transação.",
        "A exclusão conta vínculos antes de excluir; um novo vínculo concorrente pode surgir entre as duas operações.",
    ])
    add_heading(doc, "11.4 Disponibilidade e recuperação", 2)
    add_paragraph(doc, "Para produção, recomenda-se replica set com maioria de escrita adequada ao requisito, backups testados, restauração periódica, monitoramento de replication lag e política de RPO/RTO. O health check atual confirma conectividade, mas não mede capacidade de escrita, saturação do pool ou atraso de replicação.")


def chapter_12(doc):
    add_heading(doc, "12. Qualidade, testes e observabilidade", 1)
    add_heading(doc, "12.1 Situação verificada", 2)
    add_table(doc, ["Verificação", "Resultado", "Interpretação"], [
        ("go test ./...", "Aprovado", "Compilação dos pacotes; não há arquivos _test.go no repositório analisado."),
        ("npm run build", "Aprovado", "TypeScript e build Vite concluídos; bundle principal gera alerta acima de 500 kB."),
        ("Testes de integração", "Não identificados", "Fluxos MongoDB/API ainda dependem de teste manual."),
        ("Benchmark", "Não executado", "Protocolo definido no Capítulo 9."),
        ("Logs", "Parcial", "Gin e Zap presentes; sem métricas e tracing distribuído."),
    ], [2300, 1700, 5360])
    add_heading(doc, "12.2 Estratégia recomendada", 2)
    add_bullets(doc, [
        "Testes unitários dos serviços com repositórios fake para regras de status, titularidade e validação.",
        "Testes de integração com MongoDB isolado para índices únicos, filtros e operações concorrentes.",
        "Testes de API para autenticação, papéis, códigos HTTP e contratos JSON.",
        "Testes end-to-end das jornadas de cliente, lote, ativação e recarga.",
        "Benchmark automatizado versionado e executado em ambiente controlado.",
        "Métricas RED: taxa, erros e duração; métricas MongoDB de conexões, opcounters e cache.",
    ])
    add_heading(doc, "12.3 Observabilidade", 2)
    add_paragraph(doc, "A evolução indicada inclui request_id, logs estruturados sem dados sensíveis, histogramas de latência por rota, contagem de erros por operação, tracing entre handler e MongoDB e alertas para p95, taxa de erro, pool esgotado e operações lentas. O profiler do MongoDB deve ser usado com cautela e política de retenção.")
    add_figure_placeholder(doc, 11, "Evidência dos testes de build e execução")


def chapter_13(doc):
    add_heading(doc, "13. Implantação e operação", 1)
    add_heading(doc, "13.1 Variáveis de ambiente", 2)
    add_table(doc, ["Variável", "Exemplo", "Finalidade"], [
        ("APP_PORT", "8080", "Porta HTTP da API."),
        ("APP_ENV", "development", "Modo de execução e logging."),
        ("MONGO_URI", "mongodb://localhost:27017", "String de conexão."),
        ("MONGO_DB", "mvno_db", "Nome do database."),
        ("JWT_SECRET", "[segredo forte]", "Assinatura dos tokens."),
        ("JWT_EXPIRY_HOURS", "24", "Validade do JWT."),
    ], [2200, 3100, 4060])
    add_heading(doc, "13.2 Execução local", 2)
    add_numbered(doc, [
        "Iniciar uma instância MongoDB em localhost:27017.",
        "Copiar e ajustar as variáveis do arquivo .env.example.",
        "Na raiz, executar go run . para iniciar a API em http://localhost:8080.",
        "Em frontend, executar npm install e npm run dev.",
        "Acessar http://localhost:5173 e validar /health.",
    ])
    add_code(doc, '''# Backend
go run .

# Frontend
cd frontend
npm install
npm run dev''')
    add_heading(doc, "13.3 Inicialização e índices", 2)
    add_paragraph(doc, "Na subida, o sistema conecta, realiza ping, migra documentos legados de usuários e chama EnsureIndexes. Uma falha na criação de índices encerra a API, comportamento adequado para evitar operação sem as restrições esperadas. Em bases existentes, a criação de índice único pode falhar caso existam duplicidades; a migração deve prever diagnóstico e saneamento.")
    add_figure_placeholder(doc, 12, "Tela principal do sistema em execução")


def chapter_14(doc):
    add_heading(doc, "14. Recomendações priorizadas", 1)
    rows = [
        ("P0", "Transações ou idempotência", "Evitar estados parciais em cliente, ativação e lote", "Alta"),
        ("P0", "Paginação e limites", "Reduzir memória, tráfego e latência de listagens", "Alta"),
        ("P0", "Segredo JWT obrigatório", "Eliminar configuração insegura em produção", "Baixa"),
        ("P1", "Benchmarks com explain", "Validar índices e criar baseline", "Média"),
        ("P1", "Índices de ordenação", "Eliminar sorts em memória comprovados", "Baixa"),
        ("P1", "Endpoint agregado do dashboard", "Evitar baixar todos os chips", "Média"),
        ("P1", "Testes automatizados", "Proteger regras e contratos", "Média"),
        ("P1", "Validação de esquema", "Conter documentos inválidos fora da API", "Média"),
        ("P2", "Observabilidade", "Medir p95, erros e pool MongoDB", "Média"),
        ("P2", "Retenção de auditoria", "Controlar crescimento e exposição de PII", "Média"),
        ("P2", "Code splitting no frontend", "Reduzir bundle inicial acima de 500 kB", "Baixa"),
    ]
    add_table(doc, ["Prioridade", "Ação", "Benefício", "Esforço"], rows, [1300, 2450, 4010, 1600])
    add_heading(doc, "14.1 Roadmap sugerido", 2)
    add_numbered(doc, [
        "Instrumentar consultas e registrar baseline com dados representativos.",
        "Adicionar paginação, projeções e endpoint agregado de indicadores.",
        "Aplicar índices apenas após comparar planos de execução.",
        "Garantir atomicidade ou idempotência dos fluxos multicoletânea.",
        "Criar suíte de testes e pipeline de integração contínua.",
        "Formalizar segurança, backup, retenção e observabilidade para produção.",
    ])


def chapter_15(doc):
    add_heading(doc, "15. Conclusão", 1)
    add_paragraph(doc, "O sistema MVNO apresenta uma base arquitetural coerente para um projeto acadêmico e uma aplicação administrativa de pequeno a médio porte. A separação em camadas torna as responsabilidades compreensíveis, o modelo documental aproveita incorporação e snapshots de forma pertinente, e os índices únicos protegem identificadores críticos. Os históricos foram separados das entidades operacionais, evitando crescimento ilimitado do documento de chip.")
    add_paragraph(doc, "Sob a perspectiva de desempenho, os principais ganhos futuros não dependem imediatamente de infraestrutura maior, mas de consultas mais econômicas: paginação, projeções, agregações no servidor e índices alinhados às ordenações observadas. Sob a perspectiva de confiabilidade, o ponto mais importante é tratar as operações que alteram múltiplas coleções como unidades atômicas ou idempotentes.")
    add_paragraph(doc, "O relatório fornece um protocolo de medição para transformar hipóteses em evidências. Com a inclusão dos prints e resultados reservados, o trabalho poderá demonstrar não apenas a arquitetura implementada, mas também o método científico utilizado para avaliar e otimizar o banco de dados.")


def references(doc):
    add_heading(doc, "Referências", 1)
    refs = [
        "MONGODB, INC. MongoDB Manual: Indexes. Disponível em: https://www.mongodb.com/docs/manual/indexes/. Acesso em: 20 jun. 2026.",
        "MONGODB, INC. MongoDB Manual: Transactions. Disponível em: https://www.mongodb.com/docs/manual/core/transactions/. Acesso em: 20 jun. 2026.",
        "MONGODB, INC. MongoDB Manual: Explain Results. Disponível em: https://www.mongodb.com/docs/manual/reference/explain-results/. Acesso em: 20 jun. 2026.",
        "MONGODB, INC. MongoDB Go Driver Documentation. Disponível em: https://www.mongodb.com/docs/drivers/go/current/. Acesso em: 20 jun. 2026.",
        "GIN WEB FRAMEWORK. Documentation. Disponível em: https://gin-gonic.com/docs/. Acesso em: 20 jun. 2026.",
        "GO AUTHORS. The Go Programming Language Documentation. Disponível em: https://go.dev/doc/. Acesso em: 20 jun. 2026.",
        "REACT TEAM. React Documentation. Disponível em: https://react.dev/. Acesso em: 20 jun. 2026.",
        "TANSTACK. TanStack Query Documentation. Disponível em: https://tanstack.com/query/latest. Acesso em: 20 jun. 2026.",
        "OPENJS FOUNDATION. JSON Web Token Introduction. Disponível em: https://jwt.io/introduction. Acesso em: 20 jun. 2026.",
        "Projeto MVNO_ArqDB. Código-fonte local analisado em 20 jun. 2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(8)
        set_font(p.add_run(ref), size=10)


def appendices(doc):
    add_heading(doc, "Apêndice A - Catálogo resumido da API", 1)
    endpoints = [
        ("POST", "/auth/login", "Público", "Autenticar e emitir JWT"),
        ("GET", "/health", "Público", "Verificar API e MongoDB"),
        ("GET", "/api/v1/auth/me", "Autenticado", "Identidade corrente"),
        ("GET/POST", "/api/v1/clientes", "Leitura / admin,operator", "Listar e cadastrar"),
        ("PUT/DELETE", "/api/v1/clientes/:id", "admin,operator / admin", "Editar e excluir"),
        ("GET/POST", "/api/v1/planos", "Leitura / admin", "Listar e cadastrar"),
        ("PUT", "/api/v1/planos/:id", "admin", "Editar plano"),
        ("GET/POST", "/api/v1/chips", "Leitura / admin,operator", "Listar e cadastrar"),
        ("GET/POST", "/api/v1/chips/lotes", "admin,operator", "Listar e criar lotes"),
        ("GET", "/api/v1/chips/:iccid", "Autenticado", "Detalhar chip"),
        ("POST", "/api/v1/chips/:iccid/ativar", "admin,operator", "Ativar chip"),
        ("GET/POST", "/api/v1/chips/:iccid/recargas", "Leitura / admin,operator", "Histórico e recarga"),
        ("GET", "/api/v1/chips/:iccid/assinaturas", "Autenticado", "Histórico de assinaturas"),
        ("CRUD parcial", "/api/v1/users", "admin", "Gestão de usuários"),
    ]
    add_table(doc, ["Método", "Rota", "Acesso", "Finalidade"], endpoints, [1500, 3700, 1900, 2260])

    add_heading(doc, "Apêndice B - Checklist de imagens", 1)
    add_table(doc, ["Figura", "Imagem a inserir", "Como obter"], [
        ("1", "Casos de uso", "Ferramenta UML ou Mermaid exportado"),
        ("2", "Contexto da arquitetura", "Diagrama frontend/API/MongoDB"),
        ("3", "Componentes internos", "Handler, Service, Repository e Model"),
        ("4", "Sequência de ativação", "Diagrama de sequência"),
        ("5", "Modelo documental", "Coleções e cardinalidades"),
        ("6", "Coleções no Compass", "Captura do database mvno_db"),
        ("7", "Índices no Compass", "Aba Indexes das coleções"),
        ("8", "Explain antes", "executionStats sem índice proposto"),
        ("9", "Explain depois", "executionStats com índice proposto"),
        ("10", "Gráfico de benchmark", "P50/P95 e throughput"),
        ("11", "Build/testes", "Terminal com comandos aprovados"),
        ("12", "Sistema", "Dashboard ou tela de chips/lotes"),
    ], [1200, 4000, 4160])

    add_heading(doc, "Apêndice C - Roteiro de validação no MongoDB Compass", 1)
    add_numbered(doc, [
        "Conectar em mongodb://localhost:27017 e abrir mvno_db.",
        "Capturar a lista de coleções e a contagem de documentos.",
        "Abrir chips e registrar um documento com lote, cliente e assinatura atual.",
        "Abrir a aba Indexes e conferir nomes, chaves e unicidade.",
        "Executar as consultas do Capítulo 9 com explain executionStats.",
        "Salvar evidências antes/depois usando a mesma massa e mesmo limite.",
        "Substituir os quadros de figura e preencher a tabela de resultados.",
    ])

    add_heading(doc, "Apêndice D - Glossário", 1)
    add_table(doc, ["Termo", "Definição"], [
        ("MVNO", "Operadora móvel virtual que utiliza infraestrutura de rede de terceiros."),
        ("ICCID", "Identificador único do cartão SIM."),
        ("IMSI", "Identidade internacional do assinante móvel."),
        ("MSISDN", "Número telefônico associado à linha móvel."),
        ("BSON", "Formato binário de documentos utilizado pelo MongoDB."),
        ("Snapshot", "Cópia dos dados relevantes em determinado instante."),
        ("IXSCAN", "Estágio de execução que percorre um índice."),
        ("COLLSCAN", "Varredura completa da coleção."),
        ("P95", "Percentil em que 95% das medições são menores ou iguais ao valor."),
        ("RPO/RTO", "Objetivos de perda tolerável de dados e tempo de recuperação."),
    ], [2200, 7160])


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_summary(doc)
    add_toc(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    chapter_8(doc)
    chapter_9(doc)
    chapter_10(doc)
    chapter_11(doc)
    chapter_12(doc)
    chapter_13(doc)
    chapter_14(doc)
    chapter_15(doc)
    references(doc)
    appendices(doc)

    core = doc.core_properties
    core.title = "Relatório Técnico - Arquitetura e Desempenho de Banco de Dados - Sistema MVNO"
    core.subject = "Análise técnica da arquitetura Go, React e MongoDB"
    core.author = "[NOME DO(A) ALUNO(A)]"
    core.keywords = "MongoDB, MVNO, arquitetura, desempenho, índices, Go, React"
    core.comments = "Relatório gerado a partir da análise do código-fonte do projeto MVNO_ArqDB."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
